from io import BytesIO
import re
import cv2
import numpy as np
from typing import List, Dict
from loguru import logger
from PIL import Image
import pytesseract
from paddleocr import PaddleOCR
import zxingcpp
import pybit7z
import uuid
import docx
import fitz
import traceback
import zipfile
import base64
from bs4 import BeautifulSoup
from spire.doc import *
from spire.doc.common import *
from .dt import Node, File, Data
from .types import Types
import torch
from transformers import TrOCRProcessor, VisionEncoderDecoderModel
import logging
import os


os.environ["TOKENIZERS_PARALLELISM"] = "false"  
logging.getLogger("transformers").setLevel(logging.ERROR)
logging.getLogger("transformers.modeling_utils").setLevel(logging.ERROR)
logging.getLogger("transformers.configuration_utils").setLevel(logging.ERROR)
logging.getLogger("PIL.TiffImagePlugin").setLevel(logging.ERROR)

# Helper functions
def encode_binary(text: str) -> bytes:
    return text.encode('utf-8')

def decode_binary(data: bytes) -> str:
    return data.decode('utf-8')

class Extractor:
    # 添加类变量来缓存TrOCR模型和处理器
    trocr_processor = None
    trocr_model = None
    gpu_available = None
    
    @staticmethod
    def _initialize_trocr():
        """初始化OCR模型和处理器（仅在第一次需要时加载）"""
        if Extractor.gpu_available is None:
            Extractor.gpu_available = torch.cuda.is_available()
        
        if Extractor.gpu_available and Extractor.trocr_processor is None and Extractor.trocr_model is None:
            try:
                logger.info("Initializing GPU OCR model (first time only)...")
                Extractor.trocr_processor = TrOCRProcessor.from_pretrained("microsoft/trocr-base-handwritten")
                Extractor.trocr_model = VisionEncoderDecoderModel.from_pretrained("microsoft/trocr-base-handwritten").to("cuda")
                logger.info("GPU OCR model initialized successfully")
                return True
            except Exception as e:
                logger.warning(f"Failed to initialize GPU OCR: {e}. Will use CPU OCR instead.")
                Extractor.gpu_available = False
                return False
        return Extractor.gpu_available and Extractor.trocr_processor is not None and Extractor.trocr_model is not None

    @staticmethod
    def extract_urls(node: Node) -> List[Node]:
        nodes = []
        text = ""
        
        try:
            if isinstance(node.content, File):
                logger.debug(f"Node[{node.id}] file {node.content.mime_type}")
                text = decode_binary(node.content.content)
            elif isinstance(node.content, Data):
                logger.debug(f"Node[{node.id}] data {node.content.type}")
                text = decode_binary(node.content.content)
            
            urls = Extractor.extract_urls_from_text(text)
            logger.debug(f"Node[{node.id}] Number of urls: {len(urls)}")
            
            for url in urls:
                t_node = Node()
                t_node.id = 0
                t_node.content = Data(type="URL", content=encode_binary(url))
                t_node.prev = node
                # 使用新方法继承父节点的页数限制
                t_node.inherit_limits(node)
                nodes.append(t_node)
                
        except Exception as e:
            logger.error(f"Error extracting URLs: {str(e)}")
            
        return nodes

    @staticmethod
    def extract_urls_from_text(text: str) -> List[str]:
        url_pattern = r"(https?://[^\s\"<>{}]+)"
        return re.findall(url_pattern, text)

    @staticmethod
    def extract_qrcode(node: Node) -> List[Node]:
        nodes = []
        
        try:
            if isinstance(node.content, File):
                data = node.content.content
                bytes_io = BytesIO(data)
                file_bytes = np.asarray(bytearray(bytes_io.read()), dtype=np.uint8)
                img = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
                # img = cv2.imread('test.png')
                barcodes = zxingcpp.read_barcodes(img)
                for barcode in barcodes:
                    # print('Found barcode:'
                    #     f'\n Text:    "{barcode.text}"'
                    #     f'\n Format:   {barcode.format}'
                    #     f'\n Content:  {barcode.content_type}'
                    #     f'\n Position: {barcode.position}')
                
                # reader = zxingcpp.BarCodeReader()
                # image = Image.open(BytesIO(data))
                # result = reader.decode(image)
                
                # if result and result.raw:
                    t_node = Node()
                    t_node.id = 0
                    t_node.content = Data(type="QRCODE", content=encode_binary(barcode.text))
                    t_node.prev = node
                    # 使用inherit_limits方法继承限制
                    t_node.inherit_limits(node)
                    nodes.append(t_node)
                    
            elif isinstance(node.content, Data):
                logger.debug("extract_qrcode enter Data type")
                
        except Exception as e:
            logger.error(f"Error extracting QR code: {str(e)}")
            
        return nodes

    @staticmethod
    def extract_ocr(node: Node) -> List[Node]:
        nodes = []
        
        try:
            if isinstance(node.content, File):
                data = node.content.content
                
                # 检查GPU并初始化模型（仅在首次调用时）
                has_gpu_ocr = Extractor._initialize_trocr()
                extracted_text = ""
                
                # Open the image
                image = Image.open(BytesIO(data))
                
                if has_gpu_ocr:
                    try:
                        pixel_values = Extractor.trocr_processor(images=image, return_tensors="pt").pixel_values.to("cuda")
                        generated_ids = Extractor.trocr_model.generate(pixel_values)
                        extracted_text = Extractor.trocr_processor.batch_decode(generated_ids, skip_special_tokens=True)[0]
                    except Exception as e:
                        logger.warning(f"Error when using GPU OCR with GPU: {e}. Falling back to CPU OCR.")
                        has_gpu_ocr = False
                
                if not has_gpu_ocr:
                    # Fall back to PaddleOCR on CPU
                    ocr = PaddleOCR(use_angle_cls=True, lang='ch', use_gpu=False, 
                                    model_dir="/root/.paddleocr", show_log=False)
                    
                    # Convert PIL Image to numpy array as PaddleOCR needs numpy array
                    image_np = np.array(image)
                    
                    # Check if the image is RGBA and convert to RGB
                    if image_np.shape[-1] == 4:
                        image_np = cv2.cvtColor(image_np, cv2.COLOR_RGBA2RGB)
                    
                    # Run OCR
                    result = ocr.ocr(image_np, cls=True)
                    
                    if result:
                        # Extract text from results
                        text_results = []
                        for line in result:
                            for item in line:
                                if len(item) >= 2:  # Make sure the result has the expected format
                                    text_results.append(item[1][0])  # item[1][0] contains the recognized text
                        
                        # Join all recognized text
                        extracted_text = '\n'.join(text_results)
                
                if extracted_text:
                    t_node = Node()
                    t_node.id = 0
                    t_node.content = Data(type="OCR", content=encode_binary(extracted_text))
                    t_node.prev = node
                    # 使用inherit_limits方法继承限制
                    t_node.inherit_limits(node)
                    nodes.append(t_node)
                    
            elif isinstance(node.content, Data):
                logger.debug("extract_ocr enter Data type")
                
        except Exception as e:
            logger.error(f"OCR processing failed: {str(e)}")
            
        return nodes

    @staticmethod
    def extract_text_from_html(html: str) -> str:
        soup = BeautifulSoup(html, 'html.parser')
        return soup.get_text(separator=' ', strip=True)
    
    @staticmethod
    def extract_urls_from_html(html: str) -> list:
        """
        提取 HTML 中的所有 URL，覆盖常见标签、元数据、预加载、表单、懒加载、SVG 以及内联 CSS 中的 URL。
        """
        soup = BeautifulSoup(html, 'html.parser')
        urls = set()

        # 1. 基础标签属性
        # 各标签对应的 URL 属性（部分标签有多个 URL 来源）
        tag_attr_map = {
            'a': ['href'],
            'img': ['src', 'srcset'],
            'script': ['src', 'data-main'],
            'link': ['href'],
            'iframe': ['src'],
            'video': ['src', 'poster'],
            'audio': ['src'],
            'track': ['src'],
            'form': ['action'],
            'input': ['src'],  # 针对 type="image" 的情况
            'object': ['data'],
            'embed': ['src']
        }

        for tag, attrs in tag_attr_map.items():
            for element in soup.find_all(tag):
                for attr in attrs:
                    value = element.get(attr)
                    if value:
                        if attr == 'srcset':
                            # srcset 可能形如 "img1.jpg 1x, img2.jpg 2x" 分割后提取 URL
                            for part in value.split(','):
                                candidate = part.strip().split(' ')[0].strip()
                                if candidate:
                                    urls.add(candidate)
                        else:
                            urls.add(value.strip())

        # 2. 元数据与 SEO
        for meta in soup.find_all('meta'):
            # 开放图谱 <meta property="og:image" content="url">
            if meta.get('property', '').strip().lower() == 'og:image':
                content = meta.get('content')
                if content:
                    urls.add(content.strip())
            # 页面刷新 <meta http-equiv="refresh" content="5;url=redirect_url">
            if meta.get('http-equiv', '').lower() == 'refresh':
                content = meta.get('content', '')
                # 匹配形如 "5;url=redirect_url" 的格式
                m = re.search(r'url=([^;]+)', content, flags=re.IGNORECASE)
                if m:
                    urls.add(m.group(1).strip())

        # 网站图标、DNS预解析、预加载等依靠 <link> 标签，上面已经在 tag_attr_map 中处理

        # 3. 动态内容与懒加载
        # 任意含有 data-src 属性
        for element in soup.find_all(attrs={"data-src": True}):
            data_src = element.get("data-src")
            if data_src:
                urls.add(data_src.strip())

        # 4. SVG 和特殊标签
        # <image> 标签可能使用 xlink:href 或 href 属性
        for image in soup.find_all('image'):
            xlink = image.get('xlink:href')
            if xlink:
                urls.add(xlink.strip())
            alternative = image.get('href')
            if alternative:
                urls.add(alternative.strip())

        # 5. CSS 中的 URL（内联样式以及 <style> 标签内）
        # 正则匹配 url(...) 形式, 可处理单引号、双引号或不带引号的情况
        style_pattern = re.compile(r'url\((?:\'|"|)([^\'")]+)(?:\'|"|)\)')

        # 检查所有存在 style 属性的元素
        for element in soup.find_all(style=True):
            style_text = element.get('style', '')
            matches = style_pattern.findall(style_text)
            for m in matches:
                if m:
                    urls.add(m.strip())

        # 检查 <style> 标签内部的 CSS 内容
        for style_tag in soup.find_all('style'):
            css_content = style_tag.string
            if css_content:
                matches = style_pattern.findall(css_content)
                for m in matches:
                    if m:
                        urls.add(m.strip())

        return list(urls)
    

    @staticmethod
    def extract_img_from_html(html: str) -> list:
        img_bytes_list = []
        soup_t = BeautifulSoup(html, 'html.parser')
        # 查找所有的 img 标签
        images = soup_t.find_all('img')
        
        # 遍历 img 标签，打印图片 src 属性
        for image in images:
            # print(f"\n3.image --> {image}\n")
            # 这里的img_src是：data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAZAAAAGQCAIAAA……
            # print(f"\n3.image --> {type(image)} -- \n")
            if not image.has_attr('src'):
                continue
            img_src = image['src']
            # print(f"\n4.img_src --> {img_src}\n")
            #self._logger().debug(f"find image resource - {img_src}")
            # 这个判据不知道是否正确，最起码没有base64编码的就有问题了
            if (img_src.find("base64") == -1):
                continue 
            img_arrays1 = img_src.split(';')
            
            # img_arrays1[1]中的数据是：base64,iVBORw0KGgoAAAANSUhEUgAAAZAAAAGQCAIAAA……
            img_arrays2 = img_arrays1[1].split(',')
            if len(img_arrays2) < 1:
                continue
            if img_arrays2[0] != 'base64':
                continue
            
            # 这里 img_str 中是 base64 编码的：iVBORw0KGgoAAAANSUhEUgAAAZAAAAGQCAIAAA……
            img_str = img_arrays2[1]
            
            # 这里的 img_bytes 就是图片的二进制数组，base64 解码后的
            img_bytes = base64.b64decode(img_str)

            img_bytes_list.append(img_bytes)

        return img_bytes_list


    @staticmethod
    def extract_html(node: Node) -> List[Node]:
        nodes = []
        text = ""
        
        try:
            if isinstance(node.content, File):
                logger.debug(f"Node[{node.id}] file {node.content.mime_type}")
                text = decode_binary(node.content.content)
            elif isinstance(node.content, Data):
                logger.debug(f"Node[{node.id}] data {node.content.type}")
                text = decode_binary(node.content.content)
            
            html_text = Extractor.extract_text_from_html(text)
            
            t_node = Node()
            t_node.id = 0
            t_node.content = Data(type="TEXT", content=encode_binary(html_text))
            t_node.prev = node
            # 继承父节点的页数限制
            t_node.pdf_max_pages = node.pdf_max_pages
            t_node.word_max_pages = node.word_max_pages
            nodes.append(t_node)

            html_urls = Extractor.extract_urls_from_html(text)
            for url in html_urls:
                    t_node = Node()
                    t_node.id = 0
                    t_node.content = Data(type="URL", content=encode_binary(url))
                    t_node.prev = node
                    # 继承父节点的页数限制
                    t_node.pdf_max_pages = node.pdf_max_pages
                    t_node.word_max_pages = node.word_max_pages
                    nodes.append(t_node)

            img_bytes_list = Extractor.extract_img_from_html(text)
            for img_bytes in img_bytes_list:
                t_node = Node()
                t_node.content = File(
                    path="",
                    name="",
                    content=img_bytes
                )
                t_node.prev = node
                # 继承父节点的页数限制
                t_node.pdf_max_pages = node.pdf_max_pages
                t_node.word_max_pages = node.word_max_pages
                nodes.append(t_node)
            
        except Exception as e:
            logger.error(f"Error extracting HTML: {str(e)}")
            
        return nodes

    @staticmethod
    def extract_compressed_file(node: Node) -> List[Node]:
        nodes = []
        data = None
        
        try:
            if isinstance(node.content, File):
                data = node.content.content
            elif isinstance(node.content, Data):
                logger.debug("extract_compressed_file enter Data type")
                return nodes
            
            extracted = False
            files = {}
            
            if not node.passwords:
                try:
                    files = Extractor.extract_files_from_data(data)
                    extracted = True
                except Exception as e:
                    # if "password is required" not in str(e).lower():
                    # 需要密码的报错
                    # Failed to extract the archive: A password is required but none was provided.
                    raise e

            if not extracted and node.passwords:
                for password in node.passwords:
                    try:
                        files = Extractor.extract_files_from_data(data, password)
                        extracted = True
                        node.meta.map_string["correct_password"] = password
                        break
                    except Exception as e:
                        if "Wrong password" in str(e):
                            logger.error(f"Password error: {e}")
                            continue
                        raise e

            if not extracted:
                raise RuntimeError("Failed to extract compressed file")

            for filename, content in files.items():
                t_node = Node()
                t_node.content = File(
                    path=filename,
                    name=filename,
                    content=content
                )
                t_node.prev = node
                # 使用inherit_limits方法继承限制
                t_node.inherit_limits(node)
                nodes.append(t_node)

        except Exception as e:
            logger.error(f"Error extracting compressed file: {str(e)}")
            raise e

        return nodes

    @staticmethod
    def extract_files_from_data(data: bytes, password: str = "") -> Dict[str, bytes]:
        files_map = {}
        
        try:
            with pybit7z.lib7zip_context() as lib:
                extractor = pybit7z.BitMemExtractor(lib, pybit7z.FormatAuto)
                if password:
                    extractor.set_password(password)
                files_map = extractor.extract(data)
                                
        except Exception as e:
            logger.error(f"Error in extract_files_from_data: {str(e)}")
            raise e
            
        return files_map

    @staticmethod
    def extract_word_file(node: Node) -> List[Node]:
        node.meta.map_bool["is_encrypted"] = False
        nodes = []
        file: File
        if isinstance(node.content, File):
            file = node.content
        elif isinstance(node.content, Data):
            logger.error("extract_word_file enter Data type")
            return nodes
        
        def analyze_docx_file(analyzed_path:str):
            ole_file_count = 0
            doc = docx.Document(analyzed_path)
            text_content = []
            
            # 限制处理的段落数量，根据页数估算
            # 假设每页约有20个段落
            max_paragraphs = node.word_max_pages * 20
            for i, para in enumerate(doc.paragraphs):
                if i >= max_paragraphs:
                    break
                text_content.append(para.text)

            separator = "\n"
            joined_text_content = separator.join(text_content)
            # 暂时保留下面的话 👇
            # if '此文件仅WPS能打开' not in text_content:
            t_node = Node()
            t_node.id = 0
            t_node.content = Data(type="TEXT", content=encode_binary(joined_text_content))
            t_node.prev = node
            # 使用inherit_limits方法继承限制
            t_node.inherit_limits(node)
            nodes.append(t_node)

            with zipfile.ZipFile(analyzed_path, 'r') as docx_zip:
                all_files = docx_zip.namelist()
                for file1 in all_files:
                    if file1.startswith('word/embeddings/') and not file1.__eq__('word/embeddings/'):
                        ole_file_count += 1
                    elif file1.startswith('word/media/') and not file1.__eq__('word/media/'):
                        t_node = Node()
                        file_name = os.path.basename(file1)
                        file_byte = docx_zip.read(file1)
                        t_node.content = File(
                            path=file_name,
                            name=file_name,
                            content=file_byte
                        )
                        t_node.prev = node
                        # 使用inherit_limits方法继承限制
                        t_node.inherit_limits(node)
                        nodes.append(t_node)

            if ole_file_count:
                doc = Document()
                doc.LoadFromFile(analyzed_path)

                i = 1 
                for k in range(doc.Sections.Count):
                    sec = doc.Sections.get_Item(k)
                    # Iterate through all child objects in the body of each section
                    for j in range(sec.Body.ChildObjects.Count):
                        obj = sec.Body.ChildObjects.get_Item(j)
                        # Check if the child object is a paragraph
                        if isinstance(obj, Paragraph):
                            par = obj if isinstance(obj, Paragraph) else None
                            # Iterate through the child objects in the paragraph
                            for m in range(par.ChildObjects.Count):
                                o = par.ChildObjects.get_Item(m)
                                # Check if the child object is an OLE object
                                if o.DocumentObjectType == DocumentObjectType.OleObject:
                                    ole = o if isinstance(o, DocOleObject) else None
                                    s = ole.ObjectType
                                    scanners = []
                                    # Check if the OLE object is a PDF file
                                    if s.startswith("AcroExch.Document"):
                                        ext = ".pdf"
                                        scanners.append("ScanPdf")
                                    # Check if the OLE object is an Excel spreadsheet
                                    elif s.startswith("Excel.Sheet"):
                                        ext = ".xlsx"
                                    # Check if the OLE object is a PowerPoint presentation
                                    elif s.startswith("PowerPoint.Show"):
                                        ext = ".pptx"
                                    elif s.startswith("Package"):
                                        ext = ""
                                        scanners.append("ScanCompressed")
                                    elif s.startswith("Word.Document.12"):
                                        ext = ".docx"
                                        scanners.append("ScanDocx")
                                    elif s.startswith("Word.Document.8"):
                                        ext = ".doc"
                                        scanners.append("ScanDoc")
                                    else:
                                        continue
                                    t_node = Node()
                                    t_node.content = File(
                                        path=f"Output/OLE{i}{ext}",
                                        name=f"Output/OLE{i}{ext}",
                                        content=ole.NativeData
                                    )
                                    t_node.prev = node
                                    # 使用inherit_limits方法继承限制
                                    t_node.inherit_limits(node)
                                    nodes.append(t_node)                
                                    i += 1

                doc.Close()
        
        tmp_files = []

        current_file_path = ''
        file_path = '/tmp/' + uuid.uuid4().__str__()
        with open(file_path, 'wb') as f:
            f.write(file.content)
        tmp_files.append(file_path)
        current_file_path = file_path
        if node.type == Types.DOC:
            docx_path = file_path + 'x'
            document = Document()
            document.LoadFromFile(file_path)
            document.SaveToFile(docx_path, FileFormat.Docx)
            tmp_files.append(docx_path)
            current_file_path = docx_path

        try:
            analyze_docx_file(current_file_path)
        except zipfile.BadZipFile:
            node.meta.map_bool["is_encrypted"] = True
            stop = False
            for pwd in node.passwords:
                if stop: break
                try:
                    document = Document()
                    document.LoadFromFile(current_file_path, FileFormat.Docx, pwd)
                    document.RemoveEncryption()
                    document.SaveToFile(current_file_path)
                    stop = True
                except Exception as e:
                    pass
            if stop:
                analyze_docx_file(current_file_path)
        except Exception as e:
            traceback.print_exc()
        finally:
            for i in tmp_files:
                if os.path.exists(i):
                    os.remove(i)
        
        return nodes



            # for rel in doc.part.rels.values():
            #     if "image" in rel.reltype:
            #         image_path = rel.target_part.partname
            #         images.append(image_path)

            # print(text_content)
            # print(images)
            # if '此文件仅WPS能打开' not in text_content:
            #     file_name = ''
            #     if isinstance(file.original_filename, bytes):
            #         file_name = file.original_filename.decode()
            #     else:
            #         file_name = file.original_filename

            #     self.texts.append({file_name: ''.join(text_content)})
            #     extractor = URIExtractor()
            #     separator = "\n"
            #     joined_text_content = separator.join(text_content)

            #     uris:list = extractor.extracturis(' '.join(text_content))
            #     uris += extract_urls(' '.join(text_content))
            #     if set(uris):
            #         for i in uris:
            #             tmp_dict = {}
            #             tmp_dict["fileid"] = file.uid
            #             tmp_dict["string"] = i
            #             tmp_dict["flavors"] = copy.deepcopy(file.file_types)
            #             tmp_dict["from"] = "WORD-TEXT"

            #             self.urls.append(tmp_dict)
        
    @staticmethod
    def extract_pdf_file(node: Node) -> List[Node]:
        nodes = []
        file: File
        if isinstance(node.content, File):
            file = node.content
        elif isinstance(node.content, Data):
            logger.error("extract_pdf_file enter Data type")
            return nodes
        
        all_text = ""
        pdf = fitz.open(stream=BytesIO(file.content), filetype="pdf")
        if pdf.needs_pass:
            node.meta.map_bool["is_encrypted"] = True
            password_success = False
            for password in node.passwords:
                if pdf.authenticate(password):
                    password_success = True
                    node.meta.map_string["correct_password"] = password
                    break
    
            if not password_success:
                raise ValueError("PDF all passwords are invalid.")
        else:
            node.meta.map_bool["is_encrypted"] = False
            
        # 使用node.pdf_max_pages来限制处理的页数
        max_pages = min(node.pdf_max_pages, len(pdf))
        for page_number in range(max_pages):
            page = pdf.load_page(page_number)
            images = page.get_images(full=True)
            text = page.get_text()
            all_text += text
            # print(images)

            for img_index, img in enumerate(images):
                xref = img[0]
                base_image = pdf.extract_image(xref)
                image_bytes = base_image["image"]
                # image = Image.open(io.BytesIO(image_bytes))
                image_filename = f"page_{page_number + 1}_image_{img_index + 1}.png"
                t_node = Node()
                t_node.content = File(
                    path=image_filename,
                    name=image_filename,
                    content=image_bytes
                )
                t_node.prev = node
                # 使用inherit_limits方法继承限制
                t_node.inherit_limits(node)
                nodes.append(t_node)

        t_node = Node()
        t_node.id = 0
        t_node.content = Data(type="TEXT", content=encode_binary(all_text))
        t_node.prev = node
        # 使用inherit_limits方法继承限制
        t_node.inherit_limits(node)
        nodes.append(t_node)

        return nodes
